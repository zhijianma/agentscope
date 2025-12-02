# -*- coding: utf-8 -*-
# pylint: disable=W0212
"""The Word reader to read and chunk Word documents."""
import base64
import hashlib
import json
from typing import Literal, TYPE_CHECKING


from ._reader_base import ReaderBase
from ._text_reader import TextReader
from .._document import Document, DocMetadata
from ..._logging import logger
from ...message import ImageBlock, Base64Source, TextBlock

if TYPE_CHECKING:
    from docx.table import Table as DocxTable
    from docx.text.paragraph import Paragraph as DocxParagraph
else:
    DocxTable = "docx.table.Table"
    DocxParagraph = "docx.text.paragraph.Paragraph"


def _extract_text_from_paragraph(para: DocxParagraph) -> str:
    """Extract text from a paragraph, including text in text boxes and shapes.

    Args:
        para (`Paragraph`):
            The paragraph object from which to extract text.


    Returns:
        `str`:
            Extracted text
    """
    text = ""

    # Method 1: Extract all w:t elements directly from XML
    #  (handles revisions, hyperlinks, etc.)
    from docx.oxml.ns import qn

    for t_elem in para._element.findall(".//" + qn("w:t")):
        if t_elem.text:
            text += t_elem.text

    # Method 2: If no text found, try standard text property
    if not text:
        text = para.text.strip()

    # Method 3: If still no text, try to extract from text boxes and shapes
    if not text:
        # Check for text boxes (txbxContent)
        txbx_contents = para._element.findall(".//" + qn("w:txbxContent"))
        for txbx in txbx_contents:
            # Extract all text from paragraphs within the text box
            for p_elem in txbx.findall(".//" + qn("w:p")):
                for t_elem in p_elem.findall(".//" + qn("w:t")):
                    if t_elem.text:
                        text += t_elem.text

        # Check for VML text boxes - use full namespace URI
        vml_ns = "{urn:schemas-microsoft-com:vml}"
        vml_textboxes = para._element.findall(".//" + vml_ns + "textbox")
        for vml_tb in vml_textboxes:
            for p_elem in vml_tb.findall(".//" + qn("w:p")):
                for t_elem in p_elem.findall(".//" + qn("w:t")):
                    if t_elem.text:
                        text += t_elem.text

    return text.strip()


def _extract_table_data(table: DocxTable) -> list[list[str]]:
    """Extract table data, handling merged cells and preserving line breaks
    within cells.

    Args:
        table (`Table`):
            The table object from which to extract data.

    Returns:
        `list[list[str]]`:
            Table data represented as a 2D list.
    """

    from docx.oxml.ns import qn

    table_data = []
    # Extract table cell elements directly from XML
    for tr in table._element.findall(qn("w:tr")):
        row_data = []

        tcs = tr.findall(qn("w:tc"))
        for tc in tcs:
            # Extract paragraphs within the table cell (preserve line breaks)
            paragraphs = []
            for p_elem in tc.findall(qn("w:p")):
                # Obtain all text elements within the paragraph
                texts = []
                for t_elem in p_elem.findall(".//" + qn("w:t")):
                    if t_elem.text:
                        texts.append(t_elem.text)

                para_text = "".join(texts)
                if para_text:
                    # Only add non-empty paragraphs
                    paragraphs.append(para_text)

            # Use \n to join multiple paragraphs
            cell_text = "\n".join(paragraphs)
            row_data.append(cell_text)

        table_data.append(row_data)

    return table_data


def _extract_image_data(para: DocxParagraph) -> list[ImageBlock]:
    """Extract image data from a paragraph.

    Args:
        para (`Paragraph`):
            The paragraph object from which to extract images.

    Returns:
        `list[ImageBlock]`:
            A list of image blocks with base64-encoded image data
    """
    images = []

    from docx.oxml.ns import qn

    # Method 1: Find all drawing elements (modern Word format)
    drawings = para._element.findall(".//" + qn("w:drawing"))

    for drawing in drawings:
        # Try to find blip elements (embedded images)
        blips = drawing.findall(".//" + qn("a:blip"))

        for blip in blips:
            # Get the relationship ID
            embed = blip.get(qn("r:embed"))

            if embed:
                try:
                    # Get the image part from the document
                    image_part = para.part.related_parts[embed]
                    # Get the image binary data
                    image_data = image_part.blob
                    # Encode to base64
                    image_base64 = base64.b64encode(image_data).decode("utf-8")

                    # Get image format from content type
                    content_type = image_part.content_type

                    images.append(
                        ImageBlock(
                            type="image",
                            source=Base64Source(
                                type="base64",
                                data=image_base64,
                                media_type=content_type,
                            ),
                        ),
                    )
                except Exception as e:
                    logger.error(
                        "Failed to extract image: %s",
                        e,
                    )

    # Method 2: Check for pict elements (older Word format)
    picts = para._element.findall(".//" + qn("w:pict"))

    for pict in picts:
        imagedatas = pict.findall(".//" + qn("v:imagedata"))

        for imagedata in imagedatas:
            rel_id = imagedata.get(qn("r:id"))

            if rel_id:
                try:
                    image_part = para.part.related_parts[rel_id]
                    image_data = image_part.blob
                    image_base64 = base64.b64encode(image_data).decode("utf-8")

                    images.append(
                        ImageBlock(
                            type="image",
                            source=Base64Source(
                                type="base64",
                                data=image_base64,
                                media_type=image_part.content_type,
                            ),
                        ),
                    )
                except Exception as e:
                    logger.error(
                        "Failed to extract image from pict: %s",
                        e,
                    )
    return images


class WordReader(ReaderBase):
    """The reader that supports reading text, image, and table content from
    Word documents (.docx files), and chunking the text content into smaller
    pieces.

    .. note:: The table content is extracted in Markdown format.

    """

    def __init__(
        self,
        chunk_size: int = 512,
        split_by: Literal["char", "sentence", "paragraph"] = "sentence",
        include_image: bool = True,
        separate_table: bool = False,
        table_format: Literal["markdown", "json"] = "markdown",
    ) -> None:
        """Initialize the Word reader.

        Args:
            chunk_size (`int`, default to 512):
                The size of each chunk, in number of characters.
            split_by (`Literal["char", "sentence", "paragraph"]`, default to \
            "sentence"):
                The unit to split the text, can be "char", "sentence", or
                "paragraph". The "sentence" option is implemented using the
                "nltk" library, which only supports English text.
            include_image (`bool`, default to False):
                Whether to include image content in the returned document. If
                activated, the embedding model you use must support image
                input, e.g. `DashScopeMultiModalEmbedding`.
            separate_table (`bool`, default to False):
                If True, tables will be treated as a new chunk to avoid
                truncation. But note when the table exceeds the chunk size,
                it will still be truncated.
            table_format (`Literal["markdown", "json"]`, \
            default to "markdown"):
                The format to extract table content. Note if the table cell
                contains `\n`, the Markdown format may not render correctly.
                In that case, you can use the `json` format, which extracts
                the table as a JSON string of a `list[list[str]]` object.
        """
        if chunk_size <= 0:
            raise ValueError(
                f"The chunk_size must be positive, got {chunk_size}",
            )

        if split_by not in ["char", "sentence", "paragraph"]:
            raise ValueError(
                "The split_by must be one of 'char', 'sentence' or "
                f"'paragraph', got {split_by}",
            )

        if table_format not in ["markdown", "json"]:
            raise ValueError(
                "The table_format must be one of 'markdown' or 'json', "
                f"got {table_format}",
            )

        self.chunk_size = chunk_size
        self.split_by = split_by
        self.include_image = include_image
        self.separate_table = separate_table
        self.table_format = table_format

        # To avoid code duplication, we use TextReader to do the chunking.
        self._text_reader = TextReader(
            self.chunk_size,
            self.split_by,
        )

    async def __call__(
        self,
        word_path: str,
    ) -> list[Document]:
        """Read a Word document, split it into chunks, and return a list of
        Document objects. The text, image, and table content will be returned
        in the same order as they appear in the Word document.

        Args:
            word_path (`str`):
                The input Word document file path (.docx file).

        Returns:
            `list[Document]`:
                A list of Document objects, where the metadata contains the
                chunked text, doc id and chunk id.
        """

        blocks = self._get_data_blocks(word_path)

        doc_id = self.get_doc_id(word_path)
        documents = []
        for block in blocks:
            if block["type"] == "text":
                for _ in await self._text_reader(block["text"]):
                    documents.append(
                        Document(
                            metadata=DocMetadata(
                                content=_.metadata.content,
                                doc_id=doc_id,
                                # The chunk_id and total_chunks will be reset
                                chunk_id=0,
                                total_chunks=0,
                            ),
                        ),
                    )

            elif block["type"] == "image":
                documents.append(
                    Document(
                        metadata=DocMetadata(
                            content=block,
                            doc_id=doc_id,
                            chunk_id=0,
                            total_chunks=1,
                        ),
                    ),
                )

        # Set chunk ids and total chunks
        total_chunks = len(documents)
        for idx, doc in enumerate(documents):
            doc.metadata.chunk_id = idx
            doc.metadata.total_chunks = total_chunks

        return documents

    def _get_data_blocks(self, word_path: str) -> list[TextBlock | ImageBlock]:
        """This function will return a list of dicts, each dict has a
        'type' field indicating 'text', 'table', or 'image', and a
        corresponding field containing the actual data.

        Args:
            word_path (`str`):
                The input Word document file path (.docx file).

        Returns:
            `list[TextBlock | ImageBlock]`:
                A list of data blocks extracted from the Word document.
        """
        # Read the Word document
        try:
            from docx import Document as DocxDocument
            from docx.oxml import CT_P, CT_Tbl
            from docx.text.paragraph import Paragraph
            from docx.table import Table
            from docx.oxml.ns import qn

        except ImportError as e:
            raise ImportError(
                "Please install python-docx to use the Word reader. "
                "You can install it by `pip install python-docx`.",
            ) from e

        doc = DocxDocument(word_path)

        # If the last block is a table
        last_type = None

        blocks: list[TextBlock | ImageBlock] = []
        for element in doc.element.body:
            if isinstance(element, CT_P):
                para = Paragraph(element, doc)

                # Extract the text
                text = _extract_text_from_paragraph(para)

                if self.include_image:
                    # Check if the paragraph contains images
                    has_drawing = bool(
                        para._element.findall(".//" + qn("w:drawing")),
                    )
                    has_pict = bool(
                        para._element.findall(".//" + qn("w:pict")),
                    )

                    if has_drawing or has_pict:
                        # Extract the image
                        blocks.extend(_extract_image_data(para))
                        last_type = "image"

                # For current text block:
                # |   separate_table   |  True  | False  |
                # |--------------------|--------|--------|
                # | last_type == text  | append | append |
                # | last_type == image |  new   |  new   |
                # | last_type == table |  new   | append |
                # | last_type == None  |  new   |  new   |
                if (
                    last_type == "text"
                    or last_type == "table"
                    and not self.separate_table
                ):
                    blocks[-1]["text"] += "\n" + text
                else:
                    blocks.append(
                        TextBlock(
                            type="text",
                            text=text,
                        ),
                    )

                # Update last type
                last_type = "text"

            elif isinstance(element, CT_Tbl):
                # Extract the table data
                table_data = _extract_table_data(Table(element, doc))

                if self.table_format == "markdown":
                    text = self._table_to_markdown(table_data)
                else:
                    text = self._table_to_json(table_data)

                # For current table block:
                # |   separate_table   |  True  | False  |
                # |--------------------|--------|--------|
                # | last_type == text  |  new   | append |
                # | last_type == image |  new   |  new   |
                # | last_type == table |  new   | append |
                # | last_type == None  |  new   |  new   |
                if not self.separate_table and last_type in ["text", "table"]:
                    blocks[-1]["text"] += "\n" + text
                else:
                    blocks.append(
                        TextBlock(
                            type="text",
                            text=text,
                        ),
                    )

                last_type = "table"

        return blocks

    @staticmethod
    def _table_to_markdown(table_data: list[list[str]]) -> str:
        """Convert table data to Markdown format.

        Args:
            table_data (`list[list[str]]`):
                Table data represented as a 2D list.

        Returns:
            `str`:
                Table in Markdown format.
        """
        if not table_data:
            return ""

        num_cols = len(table_data[0])
        md_table = ""

        # Header row
        header_row = "| " + " | ".join(table_data[0]) + " |\n"
        md_table += header_row

        # Separator row
        separator_row = "| " + " | ".join(["---"] * num_cols) + " |\n"
        md_table += separator_row

        # Data rows
        for row in table_data[1:]:
            data_row = "| " + " | ".join(row) + " |\n"
            md_table += data_row

        return md_table

    @staticmethod
    def _table_to_json(table_data: list[list[str]]) -> str:
        """Convert table data to JSON string.

        Args:
            table_data (`list[list[str]]`):
                Table data represented as a 2D list.

        Returns:
            `str`:
                Table in JSON string format.
        """
        json_strs = [
            "<system-info>A table loaded as a JSON array:</system-info>",
        ]

        for row in table_data:
            json_strs.append(
                json.dumps(row, ensure_ascii=False),
            )

        return "\n".join(json_strs)

    def get_doc_id(self, word_path: str) -> str:
        """Generate a document ID based on the Word file path.

        Args:
            word_path (`str`):
                The Word file path.

        Returns:
            `str`:
                The generated document ID.
        """
        return hashlib.md5(word_path.encode("utf-8")).hexdigest()
